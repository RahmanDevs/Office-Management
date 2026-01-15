import os
from django.shortcuts import render
from academic.models import ExamCommittee
from django.http import JsonResponse, HttpResponse
from django.conf import settings
from docxtpl import DocxTemplate
from teachers.models import Teacher, Officers
from academic.models import ExamCommittee, DepartmentDetails, Syllabus
import uuid
from datetime import date, datetime
from django.views.decorators.csrf import csrf_exempt
import json
from core.utils import convert_number_to_words_ansi, bangla_date_format_ansi
import locale

import io
import re
import requests
import pandas as pd
from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def google_sheet_docx_page(request):
    # just render the page with input box & button
    return render(request, "exams/google_sheet_docx.html")


def generate_docx_from_google_sheet(request):
    if request.method != "POST":
        return HttpResponse("Invalid request", status=405)

    sheet_url = request.POST.get("sheet_url")
    if not sheet_url:
        return HttpResponse("No link provided", status=400)

    try:
        # Convert to CSV export link
        match = re.search(r"/d/([a-zA-Z0-9-_]+)", sheet_url)
        if not match:
            return HttpResponse("Invalid Google Sheet link", status=400)

        sheet_id = match.group(1)
        csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv"
        csv_response = requests.get(csv_url)
        csv_response.raise_for_status()

        df = pd.read_csv(io.StringIO(csv_response.text), dtype=str).fillna("")

        # Create Word document
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = "Kalpurush"  # use a Bangla-capable font
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), "Kalpurush")

        doc.add_heading("আবেদনকারীর তালিকা", 0)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ক্রমিক"
        hdr_cells[1].text = "প্রার্থীর নাম ও ঠিকানা"
        hdr_cells[2].text = "গবেষণা শিরোনাম"

        def build_info(row):
            parts = []
            if "প্রার্থীর নাম" in df.columns: parts.append(row["প্রার্থীর নাম"])
            if "পিতার নাম" in df.columns: parts.append(f"পিতা: {row['পিতার নাম']}")
            if "মাতার নাম" in df.columns: parts.append(f"মাতা: {row['মাতার নাম']}")
            addr = []
            if "গ্রাম" in df.columns: addr.append(f"গ্রাম: {row['গ্রাম']}")
            if "পোস্ট" in df.columns: addr.append(f"পোস্ট: {row['পোস্ট']}")
            if "উপজেলা" in df.columns: addr.append(f"উপজেলা: {row['উপজেলা']}")
            if "জেলা" in df.columns: addr.append(f"জেলা: {row['জেলা']}")
            if addr: parts.append(" | ".join(addr))
            if "মোবাইল" in df.columns: parts.append(f"মোবাইল: {row['মোবাইল']}")
            return " | ".join([p for p in parts if p.strip()])

        for i, row in df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(i + 1)
            cells[1].text = build_info(row)
            cells[2].text = row.get("গবেষণা শিরোনাম", "")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        response = HttpResponse(
            output.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = 'attachment; filename="applicants.docx"'
        return response

    except Exception as e:
        return HttpResponse(f"Failed: {e}", status=500)
    



def create_exam_routine(request):
    

    context = {
        'committee_head': ExamCommittee.objects.all(),
    }
    return render(request, 'exams/exam_routine.html', context)

# # make exam routine 
# def exam_routine(request):
#     return render(request, 'exams/exam_routine.html')


# from django.shortcuts import render


# # Create your views here.

def generate_moderation_letter(request):
    return HttpResponse(f"""
<ol>

""")

def generate_bill_details(request):
    # Load the template
    template_path = os.path.join(settings.BASE_DIR, 'templates/doc_file', 'Bill Details.docx')
    doc = DocxTemplate(template_path)
    exam_committee = ExamCommittee.objects.first()  # Get the first exam committee for demonstration

    # Define context (dynamic data)
    context = {
        'class': 'we.wU.AvB.Gm. ¯œvZK (m¤§vb) 3q el©',
        'semester': '2q',
        'year': exam_committee.year,
        'session': exam_committee.session,
        'chairman': exam_committee.chairman.full_name_ansi,
        'members': [member.full_name_ansi for member in exam_committee.member.all()],
        'external_member': exam_committee.external_member if exam_committee.external_member else 'N/A',  #full_name_ansi
        'date': date.today().strftime('%Y-%m-%d'),
    }
    print(context)

    # Render the template with the context
    doc.render(context)

    # Generate a unique file name and save the output
    file_name = f"generated_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(settings.MEDIA_ROOT, file_name)
    doc.save(output_path)

    # Serve the generated file as a download
    with open(output_path, 'rb') as fh:
        response = HttpResponse(fh.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        return response


@csrf_exempt  
def generate_exam_routine_docx(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            # Basic fields
            program = data.get("program")
            semester = data.get("semester")
            exam_year = data.get("exam_year")
            admission_session = data.get("admission_session")
            exam_type = data.get("exam_type")
            chair_exam_committee = Teacher.objects.get(id=data.get("chair_exam_committee"))
            courses=data.get("courses", [])
            course_rows = []
            for row in courses:
                exam_date_obj = datetime.strptime(row["exam_date"], "%Y-%m-%d")
                formatted_date = exam_date_obj.strftime("%d/%m/%Y")
                bangla_days = ['†mvgevi', 'g½jevi', 'eyaevi', 'e„n¯úwZevi', 'ïµevi', 'kwbevi', 'iweevi']
                bangla_day = bangla_days[exam_date_obj.weekday()]
                course_rows.append({
                    "course_code": row['course_code'],
                    "exam_date": formatted_date,
                    "exam_day": bangla_day,
                    "exam_session": row['exam_session'],


                })
            # Template and context
            template_path = os.path.join(settings.BASE_DIR, 'templates/doc_file', 'Exam Routine.docx')
            doc = DocxTemplate(template_path)
            context = {
                'program': program,
                'semester': semester,
                'course_rows': course_rows,
                'chair_exam_committee': chair_exam_committee.full_name_ansi,
                'exam_year': exam_year,
                'exam_type': exam_type,
                'admission_session': admission_session
            }
            # Render and save the document
            doc.render(context)
            file_name = f"generated_{uuid.uuid4()}.docx"
            output_path = os.path.join(settings.MEDIA_ROOT, file_name)
            doc.save(output_path)
            # Return the file
            with open(output_path, 'rb') as fh:
                response = HttpResponse(
                    fh.read(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename={file_name}'
                return response
            # return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    else:
        # For GET requests, render the page
        context = {
            'teachers': Teacher.objects.all(),
            'officers': Officers.objects.all()
        }
        return render(request, 'exams/exam_routine.html', context=context)
@csrf_exempt  
def generate_duty_roster_docx(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            # Basic fields
            program = data.get("program")
            semester = data.get("semester")
            exam_year = data.get("exam_year")
            exam_type = data.get("exam_type")
            chair_exam_committee = Teacher.objects.get(id=data.get("chair_exam_committee"))
            courses=data.get("courses", [])
            course_rows = []
            for row in courses:
                exam_date_obj = datetime.strptime(row["exam_date"], "%Y-%m-%d")
                formatted_date = exam_date_obj.strftime("%d/%m/%Y")
                bangla_days = ['†mvgevi', 'g½jevi', 'eyaevi', 'e„n¯úwZevi', 'ïµevi', 'kwbevi', 'iweevi']
                bangla_day = bangla_days[exam_date_obj.weekday()]
                course_rows.append({
                    "course_code": row['course_code'],
                    "exam_date": formatted_date,
                    "exam_day": bangla_day,
                    "exam_session": row['exam_session'],
                    "head_examiner": row['head_examiner'],
                    "envisilators": row['envisilators'] or [],
                    "assistants": row['assistants'] or [],
                })
            # Template and context
            template_path = os.path.join(settings.BASE_DIR, 'templates/doc_file', 'Duty Roster.docx')
            doc = DocxTemplate(template_path)
            context = {
                'program': program,
                'semester': semester,
                'course_rows': course_rows,
                'chair_exam_committee': chair_exam_committee.full_name_ansi,
                'exam_year': exam_year,
                'exam_type': exam_type
            }
            # Render and save the document
            doc.render(context)
            file_name = f"generated_{uuid.uuid4()}.docx"
            output_path = os.path.join(settings.MEDIA_ROOT, file_name)
            doc.save(output_path)
            # Return the file
            with open(output_path, 'rb') as fh:
                response = HttpResponse(
                    fh.read(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename={file_name}'
                return response
            # return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    else:
        # For GET requests, render the page
        context = {
            'teachers': Teacher.objects.all(),
            'officers': Officers.objects.all()
        }
        return render(request, 'exams/duty_roster.html', context=context)
@csrf_exempt  
def generate_exam_bill_notesheet(request):
    # Set locale to Indian (may not work on all systems)
    locale.setlocale(locale.LC_ALL, 'en_IN')
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            print(data)

            # Basic fields
            program = data.get("program")
            semester = data.get("semester")
            exam_year = data.get("exam_year")
            exam_type = data.get("exam_type")
            amount = data.get("amount")
            num_of_teacher = data.get("num_of_teacher")

            chair_exam_committee = Teacher.objects.get(id=data.get("chair_exam_committee"))

            course_rows = []
            print("Course Rows Data: ", data.get("courses", []))
            courses=data.get("courses", [])
            for row in courses:
                exam_date_obj = datetime.strptime(row["exam_date"], "%Y-%m-%d")
                formatted_date = exam_date_obj.strftime("%d/%m/%Y")
                bangla_days = ['†mvgevi', 'g½jevi', 'eyaevi', 'e„n¯úwZevi', 'ïµevi', 'kwbevi', 'iweevi']
                bangla_day = bangla_days[exam_date_obj.weekday()]

                course_rows.append({
                    "course_code": row['course_code'],
                    "exam_date": formatted_date,
                    "exam_day": bangla_day,
                    "exam_session": row['exam_session'],
                    "head_examiner": row['head_examiner'],
                    "envisilators": row['envisilators'] or [],
                    "assistants": row['assistants'] or [],
                })
            # Template and context
            template_path = os.path.join(settings.BASE_DIR, 'templates/doc_file/exam', 'exam_bill_notesheet.docx')
            doc = DocxTemplate(template_path)
            context = {
                'program': program,
                'semester': semester,
                'chair_exam_committee': chair_exam_committee.full_name_ansi,
                'exam_year': exam_year,
                'exam_type': exam_type,
                'amount': locale.format_string("%d", int(amount), grouping=True),
                'num_of_teacher': num_of_teacher,
                'amount_word': convert_number_to_words_ansi(int(amount)),
            }

            # Render and save the document
            doc.render(context)
            file_name = f"generated_{uuid.uuid4()}.docx"
            output_path = os.path.join(settings.MEDIA_ROOT, file_name)
            doc.save(output_path)

            # Return the file
            with open(output_path, 'rb') as fh:
                response = HttpResponse(
                    fh.read(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename={file_name}'
                return response
            # return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    else:
        # For GET requests, render the page
        context = {
            'teachers': Teacher.objects.all(),
            'officers': Officers.objects.all()
        }
        return render(request, 'exams/exam_bill_notesheet.html', context=context)





